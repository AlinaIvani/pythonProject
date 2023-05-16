import sys
from interface.tablehotel import *
from interface.addhotel import *
from interface.tablecountry import *
from interface.addcountry import *
from interface.tablefood import *
from interface.addfood import *
from interface.tabletourtype import *
from interface.addtype import *
from interface.tabledisc import *
from interface.adddisc import *
from check_db import *
import docx
import threading


class InterfaceHotel(QtWidgets.QWidget): #таблица отели
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form6()
        self.ui.setupUi(self)
        self.ad = AddHotel()
        self.sort = self.ui.pushButton

        self.ui.tableWidget.setSortingEnabled(True)
        self.ui.pushButton_3.clicked.connect(self.openadd)
        self.ui.pushButton_5.clicked.connect(self.delhotel)
        self.ui.pushButton_4.clicked.connect(self.cell_clicked)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton.clicked.connect(self.dbhotel)
        self.ui.pushButton_10.clicked.connect(self.print)
        self.check = self.ui.checkBox
        self.line = self.ui.lineEdit

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler2)

    def print(self): #вывод данных таблицы в текстовый документ
        rows = self.ui.tableWidget.rowCount()
        cols = self.ui.tableWidget.columnCount()
        data = []
        for row in range(rows):
            tmp = []
            for col in range(cols):
                try:
                    tmp.append(self.ui.tableWidget.item(row, col).text())
                except:
                    tmp.append('Нет данных')
            data.append(tmp)

        data_for_word = []
        for i in data:
            data_for_word.append(i)

        # добавляем таблицу
        doc = docx.Document()
        table = doc.add_table(rows=len(data_for_word), cols=4)
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for row in range(len(data_for_word)):
            for col in range(4):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(data_for_word[row][col])

        doc.save('D://отели.docx')
        QtWidgets.QMessageBox.about(self, 'Оповещение', "файл сохранен на диск D")

    def closes(self): #закрыть таблицу
        self.close()

    def dbhotel(self): #оправка данных для применения сортрировки
        threading.Timer(15, self.dbhotel).start()
        if self.ui.checkBox.isChecked() == True:
            inf1 = self.ui.lineEdit.text()
        else:
            inf1 = "%"
        if self.ui.checkBox_2.isChecked() == True:
            inf2 = self.ui.lineEdit_2.text()
        else:
            inf2 = "%"
        if self.ui.checkBox_3.isChecked() == True:
            inf3 = self.ui.lineEdit_3.text()
        else:
            inf3 = "%"
        self.check_db.thr_dbhotel(inf1, inf2, inf3)

    def signal_handler(self, value): #вывод данных в таблицу
        self.ui.tableWidget.setRowCount(len(value))
        for i, row in enumerate(value):
            item = QtWidgets.QTableWidgetItem()
            item.setData(Qt.EditRole, row[2])
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(row[0]))
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(row[1]))
            self.ui.tableWidget.setItem(i, 2, item)
            self.ui.tableWidget.setItem(i, 3, QtWidgets.QTableWidgetItem(row[3]))
            self.ui.tableWidget.resizeColumnToContents(0)
            self.ui.tableWidget.resizeColumnToContents(1)
            self.ui.tableWidget.resizeColumnToContents(2)
            self.ui.tableWidget.resizeRowsToContents()

    def delhotel(self): #удаление строки из таблицы
        row = self.ui.tableWidget.currentItem().row()
        a = self.ui.tableWidget.item(row, 0).text()
        massage = QtWidgets.QMessageBox()
        value = massage.question(self, 'Удаление', 'Вы действительно хотите удалить отель ' + a + '?', massage.Yes | massage.No)
        if value == massage.Yes:
           self.check_db.thr_delhotel(a)
        else:
           massage.close()

    def signal_handler2(self, value): #сигнал при удалении
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def cell_clicked(self): #обработка нажатий для удаления и изменение данных
        row = self.ui.tableWidget.currentItem().row()
        name = self.ui.tableWidget.item(row, 0).text()
        city = self.ui.tableWidget.item(row, 1).text()
        star = self.ui.tableWidget.item(row, 2).text()
        note = self.ui.tableWidget.item(row, 3).text()
        self.ad.nameline2.setText(name)
        self.ad.cityline.setText(city)
        self.ad.star.setEditText(star)
        self.ad.noteline.setText(note)
        self.ad.show()
        self.ad.push.hide()
        self.ad.push2.show()
        self.ad.label.hide()
        self.ad.label2.show()
        self.ad.nameline2.show()
        self.ad.nameline.hide()

    def openadd(self): #открыть окно для добавления или изменения данных
        self.ad.show()
        self.ad.push2.hide()
        self.ad.push.show()
        self.ad.label.show()
        self.ad.label2.hide()
        self.ad.nameline2.hide()
        self.ad.nameline.show()


class AddHotel(QtWidgets.QWidget): #окно добавления или изменения данных
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form01()
        self.ui.setupUi(self)

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler1)

        self.base_line_edit = [self.ui.lineEdit, self.ui.lineEdit_2]
        self.base_line_edit1 = [self.ui.lineEdit_3]
        self.nameline = self.ui.lineEdit
        self.cityline = self.ui.lineEdit_2
        self.noteline = self.ui.textEdit
        self.nameline2 = self.ui.lineEdit_3
        self.star = self.ui.comboBox
        self.push = self.ui.pushButton
        self.push2 = self.ui.pushButton_3
        self.label = self.ui.label
        self.label2 = self.ui.label_2

        self.ui.pushButton.clicked.connect(self.addhotel)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_3.clicked.connect(self.changehotel)

    def check_input(funct): #проверка правильности ввода
        def wrapper(self):
            for line_edit in self.base_line_edit:
                if len(line_edit.text()) == 0:
                    return
            funct(self)
        return wrapper

    @check_input
    def addhotel(self): #отправка данных для добавления
        name = self.ui.lineEdit.text()
        city = self.ui.lineEdit_2.text()
        star = self.ui.comboBox.currentText()
        note = self.ui.textEdit.toPlainText()
        self.check_db.thr_addhotel(name, city, star, note)

    def signal_handler(self, value): #сигнал при добавлении данных
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def changehotel(self): #отправка данных для изменения
        name = self.ui.lineEdit_3.text()
        city = self.ui.lineEdit_2.text()
        star = self.ui.comboBox.currentText()
        note = self.ui.textEdit.toPlainText()
        self.check_db.thr_changehotel(name, city, star, note)

    def signal_handler1(self, value): #сигнал при изменении данных
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def closes(self): #закрыть окно
        self.close()
        self.ui.textEdit.clear()
        self.ui.lineEdit.clear()
        self.ui.lineEdit_2.clear()


class InterfaceCountry(QtWidgets.QWidget): #таблица страны
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form7()
        self.ui.setupUi(self)
        self.ac = AddCountry()

        self.ui.tableWidget.setSortingEnabled(True)

        self.ui.pushButton.clicked.connect(self.dbcountry)
        self.ui.pushButton_3.clicked.connect(self.addopen)
        self.ui.pushButton_4.clicked.connect(self.change)
        self.ui.pushButton_5.clicked.connect(self.delcountry)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_10.clicked.connect(self.print)
        self.check = self.ui.checkBox
        self.line = self.ui.lineEdit

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler2)

    def print(self): #добавление данных из таблицы в текстовый документ
        rows = self.ui.tableWidget.rowCount()
        cols = self.ui.tableWidget.columnCount()
        data = []
        for row in range(rows):
            tmp = []
            for col in range(cols):
                try:
                    tmp.append(self.ui.tableWidget.item(row, col).text())
                except:
                    tmp.append('Нет данных')
            data.append(tmp)

        data_for_word = []
        for i in data:
            data_for_word.append(i)

        # добавляем таблицу
        doc = docx.Document()
        table = doc.add_table(rows=len(data_for_word), cols=2)
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for row in range(len(data_for_word)):
            for col in range(2):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(data_for_word[row][col])

        doc.save('D://страны.docx')
        QtWidgets.QMessageBox.about(self, 'Оповещение', "файл сохранен на диск D")

    def closes(self): #закрыть окно
        self.close()

    def delcountry(self): #удаление строки таблицы
        row = self.ui.tableWidget.currentItem().row()
        a = self.ui.tableWidget.item(row, 0).text()
        massage = QtWidgets.QMessageBox()
        value = massage.question(self, 'Удаление', 'Вы действительно хотите удалить страну ' + a + '?', massage.Yes | massage.No)
        if value == massage.Yes:
           self.check_db.thr_delcountry(a)
        else:
           massage.close()

    def signal_handler2(self, value): #сигнал при удалении
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def addopen(self):
        self.ac.show()
        self.ac.label.show()
        self.ac.label2.hide()
        self.ac.push.show()
        self.ac.push2.hide()
        self.ac.nameline.show()
        self.ac.nameline2.hide()

    def change(self):
        row = self.ui.tableWidget.currentItem().row()
        name = self.ui.tableWidget.item(row, 0).text()
        visa = self.ui.tableWidget.item(row, 1).text()
        self.ac.nameline2.setText(name)
        self.ac.visaline.setEditText(visa)
        self.ac.show()
        self.ac.label2.show()
        self.ac.label.hide()
        self.ac.push2.show()
        self.ac.push.hide()
        self.ac.nameline2.show()
        self.ac.nameline.hide()

    def dbcountry(self):
        threading.Timer(15, self.dbcountry).start()
        if self.ui.checkBox.isChecked() == True:
            inf1 = self.ui.lineEdit.text()
        else:
            inf1 = "%"
        if self.ui.checkBox_2.isChecked() == True:
            inf2 = self.ui.lineEdit_2.text()
        else:
            inf2 = "%"
        self.check_db.thr_dbcountry(inf1, inf2)

    def signal_handler(self, value):
        self.ui.tableWidget.setRowCount(len(value))
        for i, row in enumerate(value):
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(row[0]))
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(row[1]))
            self.ui.tableWidget.resizeColumnToContents(0)
            self.ui.tableWidget.resizeColumnToContents(1)

class AddCountry(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form21()
        self.ui.setupUi(self)

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler1)

        self.base_line_edit = [self.ui.lineEdit]
        self.nameline = self.ui.lineEdit
        self.visaline = self.ui.comboBox
        self.nameline2 = self.ui.lineEdit_3
        self.push = self.ui.pushButton
        self.push2 = self.ui.pushButton_3
        self.label = self.ui.label
        self.label2 = self.ui.label_2

        self.ui.pushButton.clicked.connect(self.addcountry)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_3.clicked.connect(self.changecountry)

    def check_input(funct):
        def wrapper(self):
            for line_edit in self.base_line_edit:
                if len(line_edit.text()) == 0:
                    return
            funct(self)
        return wrapper


    @check_input
    def addcountry(self):
        name = self.ui.lineEdit.text()
        visa = self.ui.comboBox.currentText()
        self.check_db.thr_addcountry(name, visa)

    def signal_handler(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def changecountry(self):
        name = self.ui.lineEdit_3.text()
        visa = self.ui.comboBox.currentText()
        print(len(self.ui.lineEdit_3.text()))
        self.check_db.thr_changecountry(name, visa)

    def signal_handler1(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def closes(self):
        self.close()
        self.ui.lineEdit.clear()

class InterfaceFood(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form8()
        self.ui.setupUi(self)
        self.af = AddFood()

        self.ui.tableWidget.setSortingEnabled(True)

        self.ui.pushButton.clicked.connect(self.dbfood)
        self.ui.pushButton_3.clicked.connect(self.addopen)
        self.ui.pushButton_4.clicked.connect(self.change)
        self.ui.pushButton_5.clicked.connect(self.delfood)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_10.clicked.connect(self.print)
        self.check = self.ui.checkBox
        self.line = self.ui.lineEdit

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler2)

    def print(self):
        rows = self.ui.tableWidget.rowCount()
        cols = self.ui.tableWidget.columnCount()
        data = []
        for row in range(rows):
            tmp = []
            for col in range(cols):
                try:
                    tmp.append(self.ui.tableWidget.item(row, col).text())
                except:
                    tmp.append('Нет данных')
            data.append(tmp)

        data_for_word = []
        for i in data:
            data_for_word.append(i)

        # добавляем таблицу
        doc = docx.Document()
        table = doc.add_table(rows=len(data_for_word), cols=2)
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for row in range(len(data_for_word)):
            for col in range(2):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(data_for_word[row][col])

        doc.save('D://питание.docx')
        QtWidgets.QMessageBox.about(self, 'Оповещение', "файл сохранен на диск D")

    def closes(self):
        self.close()

    def delfood(self):
        row = self.ui.tableWidget.currentItem().row()
        a = self.ui.tableWidget.item(row, 0).text()
        massage = QtWidgets.QMessageBox()
        value = massage.question(self, 'Удаление', 'Вы действительно хотите удалить тип питания ' + a + '?',
                                 massage.Yes | massage.No)
        if value == massage.Yes:
            self.check_db.thr_delfood(a)
        else:
            massage.close()

    def signal_handler2(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def addopen(self):
        self.af.show()
        self.af.label.show()
        self.af.label2.hide()
        self.af.push.show()
        self.af.push2.hide()
        self.af.nameline.show()
        self.af.nameline2.hide()

    def change(self):
        row = self.ui.tableWidget.currentItem().row()
        name = self.ui.tableWidget.item(row, 0).text()
        note = self.ui.tableWidget.item(row, 1).text()
        self.af.nameline2.setText(name)
        self.af.noteline.setText(note)
        self.af.show()
        self.af.label2.show()
        self.af.label.hide()
        self.af.push2.show()
        self.af.push.hide()
        self.af.nameline2.show()
        self.af.nameline.hide()

    def dbfood(self):
        threading.Timer(15, self.dbfood).start()
        if self.ui.checkBox.isChecked() == True:
            inf1 = self.ui.lineEdit.text()
        else:
            inf1 = "%"
        inf2 = "%"
        self.check_db.thr_dbfood(inf1, inf2)

    def signal_handler(self, value):
        self.ui.tableWidget.setRowCount(len(value))
        for i, row in enumerate(value):
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(row[0]))
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(row[1]))
            self.ui.tableWidget.resizeColumnToContents(0)
            self.ui.tableWidget.resizeRowsToContents()


class AddFood(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form31()
        self.ui.setupUi(self)

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler1)

        self.nameline = self.ui.lineEdit
        self.noteline = self.ui.textEdit
        self.nameline2 = self.ui.lineEdit_3
        self.push = self.ui.pushButton
        self.push2 = self.ui.pushButton_3
        self.label = self.ui.label
        self.label2 = self.ui.label_2

        self.ui.pushButton.clicked.connect(self.addfood)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_3.clicked.connect(self.changefood)

    def check_input(funct):
        def wrapper(self):
            if len(self.ui.lineEdit.text()) == 0 or len(self.ui.textEdit.toPlainText()) == 0:
                return
            funct(self)

        return wrapper

    @check_input
    def addfood(self):
        name = self.ui.lineEdit.text()
        note = self.ui.textEdit.toPlainText()
        self.check_db.thr_addfood(name, note)

    def signal_handler(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def changefood(self):
        name = self.ui.lineEdit_3.text()
        note = self.ui.textEdit.toPlainText()
        self.check_db.thr_changefood(name, note)

    def signal_handler1(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def closes(self):
        self.close()
        self.ui.lineEdit.clear()

class InterfaceType(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form9()
        self.ui.setupUi(self)
        self.at = AddType()

        self.ui.tableWidget.setSortingEnabled(True)

        self.ui.pushButton.clicked.connect(self.dbtype)
        self.ui.pushButton_3.clicked.connect(self.addopen)
        self.ui.pushButton_4.clicked.connect(self.change)
        self.ui.pushButton_5.clicked.connect(self.deltype)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_6.clicked.connect(self.print)
        self.check = self.ui.checkBox
        self.line = self.ui.lineEdit

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler2)

    def print(self):
        rows = self.ui.tableWidget.rowCount()
        cols = self.ui.tableWidget.columnCount()
        data = []
        for row in range(rows):
            tmp = []
            for col in range(cols):
                try:
                    tmp.append(self.ui.tableWidget.item(row, col).text())
                except:
                    tmp.append('Нет данных')
            data.append(tmp)

        data_for_word = []
        for i in data:
            data_for_word.append(i)

        # добавляем таблицу
        doc = docx.Document()
        table = doc.add_table(rows=len(data_for_word), cols=2)
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for row in range(len(data_for_word)):
            for col in range(2):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(data_for_word[row][col])

        doc.save('D://типы туров.docx')
        QtWidgets.QMessageBox.about(self, 'Оповещение', "файл сохранен на диск D")

    def closes(self):
        self.close()

    def deltype(self):
        row = self.ui.tableWidget.currentItem().row()
        a = self.ui.tableWidget.item(row, 0).text()
        massage = QtWidgets.QMessageBox()
        value = massage.question(self, 'Удаление', 'Вы действительно хотите удалить тип тура ' + a + '?',
                                 massage.Yes | massage.No)
        if value == massage.Yes:
            self.check_db.thr_deltype(a)
        else:
            massage.close()

    def signal_handler2(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def addopen(self):
        self.at.show()
        self.at.label.show()
        self.at.label2.hide()
        self.at.push.show()
        self.at.push2.hide()
        self.at.nameline.show()
        self.at.nameline2.hide()

    def change(self):
        row = self.ui.tableWidget.currentItem().row()
        name = self.ui.tableWidget.item(row, 0).text()
        note = self.ui.tableWidget.item(row, 1).text()
        self.at.nameline2.setText(name)
        self.at.noteline.setText(note)
        self.at.show()
        self.at.label2.show()
        self.at.label.hide()
        self.at.push2.show()
        self.at.push.hide()
        self.at.nameline2.show()
        self.at.nameline.hide()

    def dbtype(self):
        threading.Timer(15, self.dbtype).start()
        if self.ui.checkBox.isChecked() == True:
            inf1 = self.ui.lineEdit.text()
        else:
            inf1 = "%"
        inf2 = "%"
        self.check_db.thr_dbtype(inf1, inf2)

    def signal_handler(self, value):
        self.ui.tableWidget.setRowCount(len(value))
        for i, row in enumerate(value):
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(row[0]))
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(row[1]))
            self.ui.tableWidget.resizeColumnToContents(0)
            self.ui.tableWidget.resizeRowsToContents()


class AddType(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form41()
        self.ui.setupUi(self)

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler1)

        self.nameline = self.ui.lineEdit
        self.noteline = self.ui.textEdit
        self.nameline2 = self.ui.lineEdit_3
        self.push = self.ui.pushButton
        self.push2 = self.ui.pushButton_3
        self.label = self.ui.label
        self.label2 = self.ui.label_2

        self.ui.pushButton.clicked.connect(self.addtype)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_3.clicked.connect(self.changetype)

    def check_input(funct):
        def wrapper(self):
            if len(self.ui.lineEdit.text()) == 0 or len(self.ui.textEdit.toPlainText()) == 0:
                return
            funct(self)
        return wrapper

    @check_input
    def addtype(self):
        name = self.ui.lineEdit.text()
        note = self.ui.textEdit.toPlainText()
        self.check_db.thr_addtype(name, note)

    def signal_handler(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def changetype(self):
        name = self.ui.lineEdit_3.text()
        note = self.ui.textEdit.toPlainText()
        self.check_db.thr_changetype(name, note)

    def signal_handler1(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def closes(self):
        self.close()
        self.ui.lineEdit.clear()


class InterfaceDisc(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form013()
        self.ui.setupUi(self)
        self.ad = AddDisc()

        self.ui.tableWidget.setSortingEnabled(True)

        self.ui.pushButton.clicked.connect(self.dbdisc)
        self.ui.pushButton_3.clicked.connect(self.addopen)
        self.ui.pushButton_4.clicked.connect(self.change)
        self.ui.pushButton_5.clicked.connect(self.deldisc)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_10.clicked.connect(self.print)
        self.line = self.ui.lineEdit
        self.check = self.ui.checkBox

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler2)

    def print(self):
        rows = self.ui.tableWidget.rowCount()
        cols = self.ui.tableWidget.columnCount()
        data = []
        for row in range(rows):
            tmp = []
            for col in range(cols):
                try:
                    tmp.append(self.ui.tableWidget.item(row, col).text())
                except:
                    tmp.append('Нет данных')
            data.append(tmp)

        data_for_word = []
        for i in data:
            data_for_word.append(i)

        # добавляем таблицу
        doc = docx.Document()
        table = doc.add_table(rows=len(data_for_word), cols=2)
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for row in range(len(data_for_word)):
            for col in range(2):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(data_for_word[row][col])

        doc.save('D://скидки.docx')
        QtWidgets.QMessageBox.about(self, 'Оповещение', "файл сохранен на диск D")

    def closes(self):
        self.close()

    def deldisc(self):
        row = self.ui.tableWidget.currentItem().row()
        a = self.ui.tableWidget.item(row, 0).text()
        massage = QtWidgets.QMessageBox()
        value = massage.question(self, 'Удаление', 'Вы действительно хотите удалить скидку ' + a + '?',
                                 massage.Yes | massage.No)
        if value == massage.Yes:
            self.check_db.thr_deldisc(a)
        else:
            massage.close()

    def signal_handler2(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def addopen(self):
        self.ad.show()
        self.ad.label.show()
        self.ad.label2.hide()
        self.ad.push.show()
        self.ad.push2.hide()
        self.ad.nameline.show()
        self.ad.nameline2.hide()

    def change(self):
        row = self.ui.tableWidget.currentItem().row()
        name = self.ui.tableWidget.item(row, 0).text()
        note = self.ui.tableWidget.item(row, 1).text()
        self.ad.nameline2.setText(name)
        self.ad.noteline.setText(note)
        self.ad.show()
        self.ad.label2.show()
        self.ad.label.hide()
        self.ad.push2.show()
        self.ad.push.hide()
        self.ad.nameline2.show()
        self.ad.nameline.hide()

    def dbdisc(self):
        threading.Timer(15, self.dbdisc).start()
        if self.ui.checkBox.isChecked() == True:
            inf1 = self.ui.lineEdit.text()
        else:
            inf1 = "%"
        if self.ui.checkBox_2.isChecked() == True:
            inf2 = self.ui.lineEdit_2.text()
        else:
            inf2 = "%"
        self.check_db.thr_dbdisc(inf1, inf2)

    def signal_handler(self, value):
        self.ui.tableWidget.setRowCount(len(value))
        for i, row in enumerate(value):
            item = QtWidgets.QTableWidgetItem()
            item.setData(Qt.EditRole, row[1])
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(row[0]))
            self.ui.tableWidget.setItem(i, 1, item)
            self.ui.tableWidget.resizeColumnsToContents()


class AddDisc(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Form113()
        self.ui.setupUi(self)

        self.check_db = CheckThread()
        self.check_db.mysignal.connect(self.signal_handler)
        self.check_db.mysignal1.connect(self.signal_handler1)

        self.nameline = self.ui.lineEdit
        self.noteline = self.ui.lineEdit_4
        self.nameline2 = self.ui.lineEdit_3
        self.push = self.ui.pushButton
        self.push2 = self.ui.pushButton_3
        self.label = self.ui.label
        self.label2 = self.ui.label_2
        self.base_line_edit = [self.ui.lineEdit, self.ui.lineEdit_4]
        self.base_line_edit1 = [self.ui.lineEdit_3, self.ui.lineEdit_4]

        self.ui.pushButton.clicked.connect(self.adddisc)
        self.ui.pushButton_2.clicked.connect(self.closes)
        self.ui.pushButton_3.clicked.connect(self.changedisc)

    def check_input(funct):
        def wrapper(self):
            for line_edit in self.base_line_edit:
                if len(line_edit.text()) == 0:
                    return
            funct(self)
        return wrapper

    def check_input1(funct):
        def wrapper(self):
            for line_edit in self.base_line_edit1:
                if len(line_edit.text()) == 0:
                    return
            funct(self)
        return wrapper

    @check_input
    def adddisc(self):
        name = self.ui.lineEdit.text()
        note = self.ui.lineEdit_4.text()
        self.check_db.thr_adddisc(name, note)

    def signal_handler(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    @check_input1
    def changedisc(self):
        name = self.ui.lineEdit_3.text()
        note = self.ui.lineEdit_4.text()
        self.check_db.thr_changedisc(name, note)

    def signal_handler1(self, value):
        QtWidgets.QMessageBox.about(self, 'Оповещение', value)

    def closes(self):
        self.close()
        self.ui.lineEdit.clear()
        self.ui.lineEdit_4.clear()

